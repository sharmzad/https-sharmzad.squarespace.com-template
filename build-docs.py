"""Generate Word documents for Sharmzad tours and diving offerings.

Produces two .docx files with consistent brand formatting:
  - sharmzad-tours-combo.docx
  - sharmzad-diving.docx

Brand colors: gold #d4a04c, sea #1e6091, deep #0a2f4a, coral #e8743b
Contact roles:
  Diretto: WhatsApp +41 76 556 76 33 · info@egyniletravel.com
  Ufficio Egitto: WhatsApp +20 10 6209 6896 · info@sharmzad.com
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from pathlib import Path

GOLD  = RGBColor(0xD4, 0xA0, 0x4C)
CORAL = RGBColor(0xE8, 0x74, 0x3B)
SEA   = RGBColor(0x1E, 0x60, 0x91)
DEEP  = RGBColor(0x0A, 0x2F, 0x4A)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
GRAY  = RGBColor(0x6B, 0x72, 0x80)
CREAM = "FDF8F3"

def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _add_hr(paragraph, color=GOLD):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    pBdr.append(bottom)
    pPr.append(pBdr)

def brand_run(paragraph, text, size=11, color=DARK, bold=False, italic=False, uppercase=False, family="Georgia"):
    run = paragraph.add_run(text.upper() if uppercase else text)
    run.font.name = family
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        brand_run(p, text, size=22, color=DEEP, bold=True, family="Georgia")
        _add_hr(p, color=GOLD)
    elif level == 2:
        brand_run(p, text, size=15, color=SEA, bold=True, family="Georgia")
    else:
        brand_run(p, text, size=12, color=GOLD, bold=True, uppercase=True, family="Georgia")
    return p

def add_body(doc, text, size=10.5, italic=False, color=DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    brand_run(p, text, size=size, color=color, italic=italic, family="Calibri")
    return p

def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    brand_run(p, text, size=size, color=DARK, family="Calibri")

def add_cover(doc, title, subtitle, tagline):
    # Title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_run(p, "SHARMZAD", size=32, color=GOLD, bold=True, family="Georgia")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_run(p, "Tours & Diving · Sharm El Sheikh · Marsa Alam", size=9, color=GRAY, family="Calibri")

    # Title/Subtitle block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    brand_run(p, title, size=26, color=DEEP, bold=True, family="Georgia")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    brand_run(p, subtitle, size=13, color=SEA, italic=True, family="Georgia")

    # Tagline box (as 1x1 table)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, CREAM)
    cell.width = Cm(15)
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)
        brand_run(para, tagline, size=11, color=DARK, italic=True, family="Calibri")
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

def add_package(doc, name, meta, description, features):
    add_heading(doc, name, level=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    brand_run(p, meta, size=9, color=CORAL, bold=True, uppercase=True, family="Calibri")
    add_body(doc, description)
    for feat in features:
        add_bullet(doc, feat)
    p = doc.add_paragraph()
    brand_run(p, "Prezzo per persona: ", size=10, color=GRAY, family="Calibri")
    brand_run(p, "Su richiesta", size=11, color=SEA, bold=True, family="Georgia")

def add_contact_footer(doc):
    doc.add_paragraph().paragraph_format.space_before = Pt(20)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    _add_hr(p, color=GOLD)

    add_heading(doc, "Prenota ora", level=1)
    add_body(doc, "Contattaci direttamente per confermare disponibilità e finalizzare la prenotazione. Risposta entro 30 minuti. Nessun pagamento anticipato richiesto.")

    # Two-column contact table
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True

    # Direct contact (Swiss)
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, CREAM)
    for p_ in cell.paragraphs: p_.text = ""
    p = cell.paragraphs[0]
    brand_run(p, "CONTATTO DIRETTO", size=9, color=GOLD, bold=True, uppercase=True, family="Calibri")
    p = cell.add_paragraph()
    brand_run(p, "📱  +41 76 556 76 33", size=11, color=DEEP, bold=True, family="Calibri")
    p = cell.add_paragraph()
    brand_run(p, "✉️  info@egyniletravel.com", size=11, color=DEEP, family="Calibri")

    # Egypt Office
    cell = tbl.cell(0, 1)
    _set_cell_bg(cell, CREAM)
    for p_ in cell.paragraphs: p_.text = ""
    p = cell.paragraphs[0]
    brand_run(p, "UFFICIO EGITTO", size=9, color=GOLD, bold=True, uppercase=True, family="Calibri")
    p = cell.add_paragraph()
    brand_run(p, "📱  +20 10 6209 6896", size=11, color=DEEP, bold=True, family="Calibri")
    p = cell.add_paragraph()
    brand_run(p, "✉️  info@sharmzad.com", size=11, color=DEEP, family="Calibri")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_run(p, "🌐  www.sharmzad.com", size=11, color=SEA, bold=True, family="Calibri")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    brand_run(p, "© 2026 Sharmzad Tours & Diving · Sharm El Sheikh · Marsa Alam · Egitto",
              size=8, color=GRAY, italic=True, family="Calibri")

def new_document():
    doc = Document()
    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    # Base font style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc


# ============================================================
# TOURS & COMBO DOCUMENT
# ============================================================
def build_tours_doc():
    doc = new_document()
    add_cover(doc,
        title="Tour & Escursioni Sharm El Sheikh",
        subtitle="Pacchetto Combo 80€ · 6 esperienze in un unico viaggio",
        tagline="Il pacchetto più conveniente di Sharm El Sheikh — natura, deserto, cultura e mare. Guide che parlano italiano, trasferimenti dall'hotel inclusi, nessun pagamento anticipato."
    )

    # 80€ combo
    add_heading(doc, "Pacchetto Combo — 80€ a persona", level=1)
    p = doc.add_paragraph()
    brand_run(p, "TUTTO INCLUSO · 6 ESPERIENZE · 1 SOLO PREZZO",
              size=10, color=CORAL, bold=True, uppercase=True, family="Calibri")
    add_body(doc, "Dal Parco Nazionale di Ras Mohamed al tramonto sulla scogliera del Farsha Cafe. Sei esperienze diverse per vivere il Mar Rosso, il deserto del Sinai e la cultura locale in un unico pacchetto.")
    for item in [
        "Ras Mohamed via terra — Parco Nazionale del Mar Rosso, spiagge incontaminate, snorkeling tra i coralli",
        "Isola Bianca — Banco di sabbia bianca che emerge dalle acque turchesi",
        "Safari in Quad — Adrenalina tra le dune del deserto del Sinai al tramonto",
        "Cammellata nel Deserto — Esperienza beduina autentica con tè tradizionale",
        "Old Market & Città — Bazar tradizionale, moschea Al Sahaba, spezie e artigianato locale",
        "Farsha Cafe — Cuscini orientali, lanterne, vista mozzafiato sul Mar Rosso",
    ]:
        add_bullet(doc, item)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    brand_run(p, "Cosa è incluso: ", size=10.5, color=DARK, bold=True, family="Calibri")
    brand_run(p, "Tutte le 6 escursioni · trasferimenti da/per l'hotel · guide in italiano · pranzo e bevande durante le escursioni · assicurazione base · assistenza 24/7 · cancellazione gratuita fino a 24h prima.",
              size=10.5, color=DARK, family="Calibri")

    # Details per experience
    add_heading(doc, "Le sei esperienze in dettaglio", level=1)

    add_package(doc, "1 · Ras Mohamed via Terra",
        "Natura · Snorkeling · Mezza giornata",
        "Esplora il famoso Parco Nazionale di Ras Mohamed, una delle riserve marine più belle del mondo. Spiagge incantevoli, mangrovie e panorami mozzafiato sul Mar Rosso.",
        ["Spiagge incontaminate", "Snorkeling tra coralli", "Panorami unici sul Sinai"])

    add_package(doc, "2 · Isola Bianca",
        "Mare · Snorkeling · Barca",
        "Una perla nascosta del Mar Rosso: l'Isola Bianca emerge dalle acque turchesi solo con la bassa marea, regalando uno spettacolo naturale unico.",
        ["Banco di sabbia bianca", "Snorkeling tra coralli soft", "Acque cristalline"])

    add_package(doc, "3 · Safari in Quad",
        "Adrenalina · Deserto · Tramonto",
        "Vivi l'emozione di guidare un quad nel deserto del Sinai. Un'esperienza adrenalinica tra dune dorate e paesaggi lunari, con briefing in italiano.",
        ["Quad moderni e sicuri", "Tramonto nel deserto", "Briefing e guida in italiano"])

    add_package(doc, "4 · Cammellata nel Deserto",
        "Tradizione · Beduina · Serale",
        "Rivivi le antiche tradizioni beduine con una passeggiata in groppa ai cammelli. Un'esperienza autentica nel cuore del deserto del Sinai, sotto le stelle.",
        ["Cammelli docili e sicuri", "Atmosfera autentica beduina", "Tè tradizionale al campo"])

    add_package(doc, "5 · Old Market & Città",
        "Cultura · Bazar · Guida esperta",
        "Immergiti nei colori, profumi e sapori del bazar tradizionale di Sharm. Spezie, artigianato locale, moschea Al Sahaba e la vera ospitalità egiziana.",
        ["Shopping autentico", "Spezie e profumi", "Moschea Al Sahaba"])

    add_package(doc, "6 · Farsha Cafe",
        "Relax · Vista mare · Tramonto",
        "Il caffè più suggestivo di Sharm El Sheikh: arroccato sulla scogliera, con vista mozzafiato sul Mar Rosso. Atmosfera magica al tramonto, tra cuscini orientali e lanterne.",
        ["Tramonto indimenticabile", "Cuscini orientali & lanterne", "Bevande tipiche egiziane"])

    # Why us
    add_heading(doc, "Perché scegliere Sharmzad", level=1)
    for item in [
        "Guide e staff italiano — comunicazione senza barriere",
        "Miglior prezzo a Sharm — 6 esperienze a soli 80€",
        "Trasferimenti dall'hotel inclusi con veicoli climatizzati",
        "Standard di sicurezza internazionali su ogni escursione",
        "Recensioni eccellenti da migliaia di turisti italiani soddisfatti",
        "Assistenza 24/7 durante tutto il soggiorno",
    ]:
        add_bullet(doc, item)

    add_contact_footer(doc)
    out = Path("sharmzad-tours-combo.docx")
    doc.save(out)
    print(f"Wrote {out} ({out.stat().st_size:,} bytes)")


# ============================================================
# DIVING DOCUMENT
# ============================================================
def build_diving_doc():
    doc = new_document()
    add_cover(doc,
        title="Diving — Sharm & Wadi Lahami",
        subtitle="Pacchetti diving nel Mar Rosso · Sharm El Sheikh · Marsa Alam",
        tagline="Cinque pacchetti diving pronti a partire. Dai reef iconici di Sharm El Sheikh al Fury Shoals accessibile solo dall'eco village di Wadi Lahami. Guide PADI/SSI che parlano italiano."
    )

    # SHARM
    add_heading(doc, "Sharm El Sheikh — 3 pacchetti diving", level=1)
    add_body(doc, "Sharm El Sheikh — capitale mondiale del diving. Ras Mohamed National Park, Straits of Tiran, SS Thistlegorm e i reef locali. Pacchetti dal weekend lungo alla settimana completa.")

    add_package(doc, "Sharm El Sheikh — 6 giorni / 5 notti / 8 immersioni  ⭐",
        "Multi-day · Featured · Top pacchetto",
        "Il pacchetto Sharm più completo: 8 immersioni sui migliori siti dell'area — coralli spettacolari e vita marina ricchissima. Sistemazione confortevole ed esperienza personalizzata.",
        ["8 immersioni guidate", "Sistemazione 5 notti", "Immersioni sui siti top di Sharm", "Programma personalizzato"])

    add_package(doc, "Sharm El Sheikh — 5 giorni / 4 notti / 6 immersioni (privato)",
        "Private · 5 giorni · Esclusivo",
        "Pacchetto privato a Sharm El Sheikh con 6 immersioni. Sistemazione riservata ed esplorazione del Mar Rosso in massima esclusività.",
        ["6 immersioni", "Sistemazione privata 4 notti", "Esperienza esclusiva", "Programma personalizzato"])

    add_package(doc, "Sharm El Sheikh — 4 giorni / 3 notti / 4 immersioni",
        "Short getaway · Weekend lungo",
        "Getaway diving compatto: 4 immersioni, hotel 3 stelle confortevole, tempo per rilassarsi tra le uscite. Ideale per scoprire i tesori sottomarini di Sharm in un viaggio breve.",
        ["4 immersioni", "Hotel 3 stelle", "Weekend lungo", "Programma flessibile"])

    # WADI LAHAMI
    add_heading(doc, "Wadi Lahami — 2 pacchetti safari nel far south", level=1)
    add_body(doc, "Wadi Lahami Eco Village — sulla costa remota del sud dell'Egitto — dà accesso diretto al leggendario sistema di reef Fury Shoals. Deserto incontaminato, acqua cristallina, immersioni pure. In collaborazione con Repack Travel.")

    add_package(doc, "Wadi Lahami — 7 notti + 5 giorni di diving  ⭐",
        "Marsa Alam · Full South · Eco Village",
        "Sul remoto sud dell'Egitto, Wadi Lahami Eco Village offre accesso diretto al leggendario sistema di reef Fury Shoals — tra le aree più intatte e biodiverse del Mar Rosso. Deserto e acqua cristallina.",
        ["7 notti eco village", "5 giorni di immersioni", "Accesso al sistema Fury Shoals", "Deserto incontaminato", "Esperienza remota"])

    add_package(doc, "Wadi Lahami — 5 notti + 3 giorni di diving",
        "Marsa Alam · Eco Village · Intro sud",
        "Eco-diving village nel profondo sud del Mar Rosso, con accesso diretto allo spettacolare Fury Shoals — uno degli ambienti corallini più incontaminati del Mar Rosso.",
        ["5 notti eco village", "3 giorni di immersioni", "Fury Shoals system", "Ambiente naturale", "Ideale come introduzione al sud"])

    # Dive sites
    add_heading(doc, "I siti d'immersione", level=1)
    for site, desc in [
        ("Shark & Yolanda Reef (Sharm)", "Muro verticale nel cuore di Ras Mohamed. Squali di reef, tonni, barracuda + il relitto del Yolanda."),
        ("SS Thistlegorm (Sharm)", "Nave WWII con moto Norton, camion Bedford, locomotive intatte — 15-30m — AOWD+."),
        ("Jackson Reef (Sharm)", "Il più a nord degli Straits of Tiran. Pesci pelagici, tartarughe, squali martello — corrente."),
        ("Fury Shoals (Wadi Lahami)", "Sistema di reef del sud accessibile da Wadi Lahami. Giardini di coralli soft e drift dive tra pinnacoli."),
        ("Sataya Reef (Wadi Lahami)", 'Il vero "Dolphin Reef". Grandi pod di spinner dolphins in laguna coralline.'),
        ("St. John's Reefs (Far South)", "Sistema di reef nel far south, vicino al confine con il Sudan — solo su safari."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        brand_run(p, f"{site} — ", size=10.5, color=DEEP, bold=True, family="Calibri")
        brand_run(p, desc, size=10.5, color=DARK, family="Calibri")

    # Courses
    add_heading(doc, "Corsi PADI / SSI", level=1)
    for course, dur, desc in [
        ("Discover Scuba", "½ giornata", "La tua prima prova. Nessuna esperienza richiesta."),
        ("Open Water Diver", "3-4 giorni", "Primo brevetto internazionale. Fino a 18m."),
        ("Advanced Open Water", "2-3 giorni", "Cinque immersioni specialità. Fino a 30m."),
        ("Nitrox / Deep / Wreck", "1-2 giorni", "Corsi specialità per estendere tempo di fondo, profondità, siti wreck."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        brand_run(p, f"{course} · ", size=10.5, color=DEEP, bold=True, family="Calibri")
        brand_run(p, f"{dur} — ", size=10.5, color=CORAL, italic=True, family="Calibri")
        brand_run(p, desc, size=10.5, color=DARK, family="Calibri")

    # Why us
    add_heading(doc, "Perché scegliere Sharmzad", level=1)
    for item in [
        "Guide e istruttori che parlano italiano fluentemente",
        "Diving center certificati PADI e SSI, riconosciuti internazionalmente",
        "Un solo partner per Sharm + Marsa Alam + Wadi Lahami",
        "Attrezzatura completa e ben mantenuta inclusa",
        "Trasferimenti dall'hotel inclusi",
        "Assistenza 24/7 durante il soggiorno",
    ]:
        add_bullet(doc, item)

    add_contact_footer(doc)
    out = Path("sharmzad-diving.docx")
    doc.save(out)
    print(f"Wrote {out} ({out.stat().st_size:,} bytes)")


if __name__ == "__main__":
    build_tours_doc()
    build_diving_doc()
